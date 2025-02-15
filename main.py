from azure.storage.blob import BlobServiceClient, ContentSettings
import streamlit as st
import json
import redis
from io import BytesIO
from docx import Document
from urllib.parse import urlparse
from pdf_processing import process_pdf_task
from respondent import ask_question, bing_search_topics
from utils.config import (
    redis_host,
    redis_pass,
    azure_blob_connection_string,
    azure_container_name,
    bing_key,
    bing_endpoint,
)
import uuid
import tiktoken
import time
import requests

def count_tokens(text, model="gpt-4o"):
    """Count the number of tokens in the text for a specific model."""
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)


redis_client = redis.Redis(
    host=redis_host,
    port=6379,
    password=redis_pass,
)


blob_service_client = BlobServiceClient.from_connection_string(
    azure_blob_connection_string
)
container_client = blob_service_client.get_container_client(azure_container_name)
if not container_client.exists():
    container_client.create_container()


if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if "documents" not in st.session_state:
    st.session_state.documents = {}  
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "doc_token" not in st.session_state:
    st.session_state.doc_token = 0
if "removed_documents" not in st.session_state:
    st.session_state.removed_documents = []  


def save_document_to_redis(session_id, doc_id, document_data):
    """Save document data to Redis."""
    redis_key = f"{session_id}:document_data:{doc_id}"
    redis_client.set(redis_key, json.dumps(document_data))


def upload_to_blob_storage(file_name, file_data):
    """Upload a file to Azure Blob Storage."""
    try:
        blob_client = container_client.get_blob_client(file_name)
        blob_client.upload_blob(
            file_data,
            content_settings=ContentSettings(content_type="application/pdf"),
            overwrite=True,
        )
        st.success(f"Uploaded {file_name} to Azure Blob Storage!")
    except Exception as e:
        st.error(f"Error uploading to Azure Blob Storage: {e}")


import requests

def search_bing(query, bing_key, bing_endpoint):
    """Search for the top 3 Bing results."""
    # Set up headers and parameters
    headers = {"Ocp-Apim-Subscription-Key": bing_key}
    params = {"q": query, "textDecorations": True, "textFormat": "HTML", "count": 3}

    # Perform the GET request
    response = requests.get(bing_endpoint, headers=headers, params=params)

    # Raise an error if the request was not successful
    response.raise_for_status()

    # Parse the JSON response
    search_results = response.json()

    # Extract the URLs of the top 3 results
    results = []
    for web_page in search_results.get("webPages", {}).get("value", []):
        results.append(web_page["url"])

    return results



def handle_question(prompt, spinner_placeholder):
    """Handle user question by querying the documents in the session and adding Bing search results."""
    if prompt:
        try:
            documents_data = {
                doc_id: doc_info["data"]
                for doc_id, doc_info in st.session_state.documents.items()
            }
            if not documents_data:
                st.warning(
                    "No documents available in the session to answer the question."
                )
                return

            with spinner_placeholder.container():
                st.spinner("Thinking...")
                answer, tot_tokens = ask_question(
                    documents_data, prompt, st.session_state.chat_history
                )
            bing_search_query = str(f"""{prompt}\n{answer}""")
            search_str = bing_search_topics(bing_search_query)
            # Get top 3 Bing search results
            bing_results = search_bing(search_str, bing_key, bing_endpoint)

            # Add the Bing search results to the answer
            answer += "\n\nGo to the internet for more information:\n"
            for i, link in enumerate(bing_results, start=1):
                # Parse the URL and extract the hostname (domain)
                domain = urlparse(link).netloc
                answer += f"{i}. [{domain}]({link})\n"

            # Append the question and answer to chat history
            st.session_state.chat_history.append(
                {
                    "question": prompt,
                    "answer": f"{answer}\nTotal tokens: {tot_tokens}",
                }
            )
        except Exception as e:
            st.error(f"Error processing question: {e}")
        finally:
            spinner_placeholder.empty()


def display_chat():
    """Display chat history with download buttons."""
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history):
            with st.chat_message("user"):
                st.write(chat["question"])
            with st.chat_message("assistant"):
                st.write(chat["answer"])

                # Create a Word document with formatted content
                doc = Document()
                doc.add_heading("Chat Response", level=1)
                doc.add_paragraph("Question:", style="Heading 2")
                doc.add_paragraph(chat["question"])

                doc.add_paragraph("Answer:", style="Heading 2")
                
                # Format answer content
                for line in chat["answer"].split("\n"):
                    if line.startswith("#### "):  # Convert to Heading 3
                        doc.add_heading(line.replace("#### ", ""), level=3)
                    elif line.startswith("- **"):  # Bold for list items
                        line = line.replace("- **", "").replace("**:", ":")
                        doc.add_paragraph(line, style="List Bullet")
                    elif line.startswith("- "):  # Normal bullet points
                        doc.add_paragraph(line, style="List Bullet")
                    elif line.startswith("### "):  # Convert to Heading 2
                        doc.add_heading(line.replace("### ", ""), level=2)
                    elif line.strip():  # Add as normal paragraph
                        doc.add_paragraph(line)

                # Save the document to a BytesIO object
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Add a download button for the Word document
                st.download_button(
                    label="Download",
                    data=doc_io,
                    file_name=f"chat_response_{i+1}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download response as word document",
                    icon="📥"
                )



st.title("docQuest")
st.subheader("Unveil the Essence, Compare Easily, Analyze Smartly")

with st.sidebar:
    with st.expander("Document(s) are ready:", expanded=True):
        to_remove = []
        for doc_id, doc_info in st.session_state.documents.items():
            col1, col2 = st.columns([4, 1])
            with col1:
                st.write(f"{doc_info['name']}")
            with col2:
                if st.button(f"⨯", key=f"remove_{doc_id}"):
                    to_remove.append(doc_id)

        for doc_id in to_remove:
            st.session_state.doc_token -= count_tokens(
                str(st.session_state.documents[doc_id]["data"])
            )
            st.session_state.removed_documents.append(
                st.session_state.documents[doc_id]["name"]
            )
            redis_client.delete(f"{st.session_state.session_id}:document_data:{doc_id}")
            del st.session_state.documents[doc_id]
            st.success("Document removed successfully!")
            time.sleep(1)
            st.rerun()

with st.sidebar:
    with st.expander("Upload Document(s)", expanded=False):
        uploaded_files = st.file_uploader(
            "Upload files less than 400 pages",
            type=["pdf", "docx", "xlsx", "pptx", "xls"],
            accept_multiple_files=True,
            help="If your question is not answered properly or there's an error, consider uploading smaller documents or splitting larger ones.",
            label_visibility="collapsed",
        )

        if uploaded_files:
            new_files = []
            for uploaded_file in uploaded_files:
                
                if (
                    uploaded_file.name
                    not in [
                        st.session_state.documents[doc_id]["name"]
                        for doc_id in st.session_state.documents
                    ]
                    and uploaded_file.name not in st.session_state.removed_documents
                ):
                    new_files.append(uploaded_file)

            if new_files:
                progress_text = st.empty()
                progress_bar = st.progress(0)
                total_files = len(new_files)

                with st.spinner("Learning about your document(s)..."):
                    try:
                        for i, uploaded_file in enumerate(new_files):
                            document_data = process_pdf_task(
                                uploaded_file, first_file=(i == 0)
                            )
                            if not document_data:
                                st.warning(
                                    "The document exceeds the size limit for processing!",
                                    icon="⚠️",
                                )
                                uploaded_file.seek(0)
                                continue

                            doc_token_count = count_tokens(str(document_data))
                            if st.session_state.doc_token + doc_token_count > 600000:
                                st.warning(
                                    "Document contents so far are too large to query. Not processing further documents. "
                                    "Results may be inaccurate; consider uploading smaller documents.",
                                    icon="⚠️",
                                )
                                continue

                            doc_id = str(uuid.uuid4())
                            st.session_state.documents[doc_id] = {
                                "name": uploaded_file.name,
                                "data": document_data,
                            }
                            st.session_state.doc_token += doc_token_count
                            save_document_to_redis(
                                st.session_state.session_id, doc_id, document_data
                            )
                            
                            upload_to_blob_storage(
                                uploaded_file.name, uploaded_file.getvalue()
                            )
                            st.success(f"{uploaded_file.name} processed!")
                            time.sleep(1)
                            st.rerun()
                            progress_bar.progress((i + 1) / total_files)
                    except Exception as e:
                        st.error(f"Error processing file: {e}")

                progress_text.text("Processing complete.")
                progress_bar.empty()
                st.rerun()

if st.session_state.documents:
    prompt = st.chat_input("Ask me anything about your documents", key="chat_input")
    spinner_placeholder = st.empty()
    if prompt:
        handle_question(prompt, spinner_placeholder)

display_chat()
